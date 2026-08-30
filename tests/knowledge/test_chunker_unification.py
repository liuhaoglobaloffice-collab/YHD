"""P2-6 Chunker unification tests (Gap 10 + 12).

Verifies that the two chunker implementations are bridged:
- DocumentProcessor (ChunkMetadata, paragraph-based) remains unchanged
- ChunkMetadata.to_chunk() converts to the EmbeddingPipeline's Chunk type
- The unified output flows into EmbeddingPipeline.run_chunks() without
  re-chunking
- Covers TXT / PDF / DOCX / XLSX inputs, empty documents, and very long text
"""

import asyncio
import io

import pytest

from src.knowledge.chunker import Chunk, TextChunker
from src.knowledge.embedding import EmbeddingPipeline
from src.knowledge.processing import ChunkType, DocumentProcessor
from src.knowledge.vector_store import InMemoryVectorStore


def make_processor(max_chunk_size=1000) -> DocumentProcessor:
    from src.knowledge.processing import Chunker

    return DocumentProcessor(chunker=Chunker(max_chunk_size=max_chunk_size))


# ======================================================================
# to_chunk() adapter (Gap 10)
# ======================================================================


def test_to_chunk_returns_chunk_type_with_preserved_fields():
    from src.knowledge.processing import ChunkMetadata

    meta = ChunkMetadata(
        chunk_id="doc-1_c0",
        document_id="doc-1",
        document_version=1,
        chunk_index=0,
        chunk_type=ChunkType.HEADING,
        content="# Introduction",
        page=3,
        section="intro",
    )
    chunk = meta.to_chunk()

    assert isinstance(chunk, Chunk)
    assert chunk.chunk_id == "doc-1_c0"
    assert chunk.document_id == "doc-1"
    assert chunk.content == "# Introduction"
    assert chunk.chunk_index == 0
    assert chunk.metadata["chunk_type"] == "heading"
    assert chunk.metadata["document_version"] == 1
    assert chunk.metadata["page"] == 3
    assert chunk.metadata["section"] == "intro"


def test_to_chunk_omits_none_source_fields():
    from src.knowledge.processing import ChunkMetadata

    meta = ChunkMetadata(
        chunk_id="d_c0",
        document_id="d",
        document_version=1,
        chunk_index=0,
        chunk_type=ChunkType.PARAGRAPH,
        content="plain text",
    )
    chunk = meta.to_chunk()
    assert "page" not in chunk.metadata
    assert "section" not in chunk.metadata
    assert "sheet" not in chunk.metadata


def test_to_chunk_merges_custom_metadata():
    from src.knowledge.processing import ChunkMetadata

    meta = ChunkMetadata(
        chunk_id="d_c0",
        document_id="d",
        document_version=2,
        chunk_index=0,
        chunk_type=ChunkType.TABLE,
        content="a | b",
        sheet="Sheet1",
        metadata={"source": "upload", "custom": 42},
    )
    chunk = meta.to_chunk()
    assert chunk.metadata["source"] == "upload"
    assert chunk.metadata["custom"] == 42
    assert chunk.metadata["sheet"] == "Sheet1"
    assert chunk.metadata["chunk_type"] == "table"


# ======================================================================
# TXT input
# ======================================================================


@pytest.mark.asyncio
async def test_txt_document_processing_and_chunk_structure():
    processor = make_processor()
    text = "First paragraph about topics.\n\nSecond paragraph with more content."
    chunks = await processor.process_document(
        document_id="txt-1",
        document_version=1,
        file_type="text/plain",
        content=text.encode("utf-8"),
    )
    # _clean_text collapses whitespace, so the document becomes one paragraph
    assert len(chunks) == 1
    c = chunks[0]
    assert c.document_id == "txt-1"
    assert c.document_version == 1
    assert c.chunk_type == ChunkType.PARAGRAPH
    assert "First paragraph about topics." in c.content
    assert "Second paragraph with more content." in c.content
    assert c.char_count == len(c.content)

    # Unified conversion produces embeddable Chunks
    unified = [c.to_chunk() for c in chunks]
    assert all(isinstance(u, Chunk) for u in unified)


# ======================================================================
# PDF input
# ======================================================================


def _make_text_pdf() -> bytes:
    """Handcraft a tiny PDF whose single page contains extractable text."""
    text = "Hello PDF knowledge base content."
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref_pos = out.tell()
    out.write(b"xref\n0 " + str(len(objects) + 1).encode() + b"\n")
    out.write(b"0000000000 65535 f \n")
    for off in offsets:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(
        b"trailer\n<< /Size " + str(len(objects) + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return out.getvalue()


@pytest.mark.asyncio
async def test_pdf_document_processing():
    processor = make_processor()
    chunks = await processor.process_document(
        document_id="pdf-1",
        document_version=1,
        file_type="application/pdf",
        content=_make_text_pdf(),
    )
    assert len(chunks) >= 1
    combined = " ".join(c.content for c in chunks)
    assert "Hello PDF knowledge base content." in combined
    unified = [c.to_chunk() for c in chunks]
    assert all(isinstance(u, Chunk) for u in unified)


@pytest.mark.asyncio
async def test_pdf_without_extractable_text_is_marked():
    processor = make_processor()
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    chunks = await processor.process_document(
        document_id="pdf-blank",
        document_version=1,
        file_type="application/pdf",
        content=buf.getvalue(),
    )
    assert len(chunks) >= 1
    assert "[PDF Content" in chunks[0].content


# ======================================================================
# DOCX input
# ======================================================================


def _make_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("DOCX paragraph one about suppliers.")
    doc.add_paragraph("DOCX paragraph two about risk.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_docx_document_processing():
    processor = make_processor()
    chunks = await processor.process_document(
        document_id="docx-1",
        document_version=1,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=_make_docx_bytes(),
    )
    assert len(chunks) >= 1
    combined = " ".join(c.content for c in chunks)
    assert "paragraph one" in combined
    assert "paragraph two" in combined
    unified = [c.to_chunk() for c in chunks]
    assert all(isinstance(u, Chunk) for u in unified)


# ======================================================================
# XLSX input
# ======================================================================


def _make_xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Value"])
    ws.append(["alpha", "1"])
    ws.append(["beta", "2"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_xlsx_document_processing():
    processor = make_processor()
    chunks = await processor.process_document(
        document_id="xlsx-1",
        document_version=1,
        file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=_make_xlsx_bytes(),
    )
    assert len(chunks) >= 1
    combined = " ".join(c.content for c in chunks)
    assert "alpha" in combined
    assert "beta" in combined
    unified = [c.to_chunk() for c in chunks]
    assert all(isinstance(u, Chunk) for u in unified)


# ======================================================================
# Edge cases: empty document and very long text
# ======================================================================


@pytest.mark.asyncio
async def test_empty_document_returns_no_chunks():
    processor = make_processor()
    for file_type, content in [
        ("text/plain", b""),
        ("text/plain", b"   \n\n  "),
    ]:
        chunks = await processor.process_document(
            document_id="empty-1",
            document_version=1,
            file_type=file_type,
            content=content,
        )
        assert chunks == []


@pytest.mark.asyncio
async def test_very_long_text_is_split_into_multiple_chunks():
    processor = make_processor(max_chunk_size=200)
    long_paragraph = ("This is a long sentence about knowledge. " * 40).strip()
    assert len(long_paragraph) > 1000

    chunks = await processor.process_document(
        document_id="long-1",
        document_version=1,
        file_type="text/plain",
        content=long_paragraph.encode("utf-8"),
    )
    assert len(chunks) >= 2
    # Chunk indices are sequential and ids unique
    for i, c in enumerate(chunks):
        assert c.chunk_index == i
        assert c.chunk_id == f"long-1_c{i}"
    # No content is lost (whitespace normalized)
    assert sum(len(c.content) for c in chunks) >= len(long_paragraph) - 50


@pytest.mark.asyncio
async def test_pipeline_run_chunks_with_processor_output():
    """End-to-end: DocumentProcessor → to_chunk() → EmbeddingPipeline.run_chunks()."""
    processor = make_processor()
    text = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
    metas = await processor.process_document(
        document_id="flow-1",
        document_version=1,
        file_type="text/plain",
        content=text.encode("utf-8"),
    )

    store = InMemoryVectorStore()
    pipeline = EmbeddingPipeline(vector_store=store, provider_name="mock")
    result = await pipeline.run_chunks([m.to_chunk() for m in metas])

    assert result["document_id"] == "flow-1"
    assert result["chunks_count"] == len(metas)
    assert result["embeddings_created"] == len(metas)
    assert result["storage_status"] == "ok"
    # Vector store contains exactly the processed chunks (no re-chunking)
    assert {r.chunk_id for r in store.records} == {m.chunk_id for m in metas}


@pytest.mark.asyncio
async def test_run_chunks_is_idempotent_across_processor_outputs():
    processor = make_processor()
    metas = await processor.process_document(
        document_id="flow-2",
        document_version=1,
        file_type="text/plain",
        content=b"Content paragraph one.\n\nContent paragraph two.",
    )

    store = InMemoryVectorStore()
    pipeline = EmbeddingPipeline(vector_store=store, provider_name="mock")
    unified = [m.to_chunk() for m in metas]

    result1 = await pipeline.run_chunks(unified)
    result2 = await pipeline.run_chunks(unified)
    assert result1["embeddings_created"] == len(metas)
    assert result2["embeddings_skipped"] == len(metas)
    assert result2["embeddings_created"] == 0


# ======================================================================
# TextChunker (embedding-side chunker) remains compatible
# ======================================================================


def test_text_chunker_still_works_directly():
    chunker = TextChunker(chunk_size=30, overlap=5)
    chunks = chunker.chunk_text(
        document_id="tc-1",
        text="Alpha beta gamma delta epsilon zeta eta theta.",
    )
    assert len(chunks) >= 1
    assert isinstance(chunks[0], Chunk)
    assert chunks[0].chunk_id == "tc-1_chunk_0"
