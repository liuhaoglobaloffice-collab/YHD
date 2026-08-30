"""
Document Processing Pipeline

Handles document parsing, chunking, and metadata extraction.
"""

import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from enum import Enum
from typing import Any, Dict, List, Optional


class ChunkType(str, Enum):
    """Chunk type"""

    PARAGRAPH = "paragraph"
    HEADING = "heading"
    TABLE = "table"
    LIST = "list"
    CODE = "code"
    METADATA = "metadata"


@dataclass
class ChunkMetadata:
    """Chunk metadata"""

    chunk_id: str
    document_id: str
    document_version: int
    chunk_index: int
    chunk_type: ChunkType
    content: str

    # Source location
    page: Optional[int] = None
    section: Optional[str] = None
    sheet: Optional[str] = None

    # Processing metadata
    token_count: int = 0
    char_count: int = 0

    # Custom metadata
    metadata: Dict[str, Any] = field(default_factory=dict)

    created_at: datetime = field(default_factory=lambda: datetime.now(UTC))

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "document_version": self.document_version,
            "chunk_index": self.chunk_index,
            "chunk_type": self.chunk_type.value,
            "content": self.content,
            "page": self.page,
            "section": self.section,
            "sheet": self.sheet,
            "token_count": self.token_count,
            "char_count": self.char_count,
            "metadata": self.metadata,
            "created_at": self.created_at.isoformat(),
        }

    def to_chunk(self):
        """Convert to the storage-friendly Chunk used by the Embedding Pipeline.

        P2-6 chunker unification: this adapter lets chunks produced by
        DocumentProcessor flow directly into EmbeddingPipeline without
        re-chunking. Rich metadata (chunk_type, page, section, sheet) is
        preserved in the Chunk metadata dict.
        """
        from .chunker import Chunk

        meta: Dict[str, Any] = {
            "chunk_type": self.chunk_type.value,
            "document_version": self.document_version,
        }
        if self.page is not None:
            meta["page"] = self.page
        if self.section is not None:
            meta["section"] = self.section
        if self.sheet is not None:
            meta["sheet"] = self.sheet
        meta.update(self.metadata or {})

        return Chunk(
            chunk_id=self.chunk_id,
            document_id=self.document_id,
            content=self.content,
            chunk_index=self.chunk_index,
            metadata=meta,
        )


class Chunker:
    """
    Document Chunker

    Splits documents into chunks with context preservation.
    """

    # Chunk size limits
    MIN_CHUNK_SIZE = 100
    MAX_CHUNK_SIZE = 1000
    OVERLAP_SIZE = 100

    def __init__(
        self,
        min_chunk_size: int = MIN_CHUNK_SIZE,
        max_chunk_size: int = MAX_CHUNK_SIZE,
        overlap_size: int = OVERLAP_SIZE,
    ):
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size

    def chunk_text(
        self,
        document_id: str,
        document_version: int,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[ChunkMetadata]:
        """
        Chunk text content

        Args:
            document_id: Document ID
            document_version: Document version
            content: Text content
            metadata: Additional metadata

        Returns:
            List[ChunkMetadata]: List of chunks
        """
        chunks = []

        # Split by paragraphs
        paragraphs = self._split_paragraphs(content)

        chunk_index = 0
        for para_idx, paragraph in enumerate(paragraphs):
            # Skip empty paragraphs
            if not paragraph.strip():
                continue

            # Detect chunk type
            chunk_type = self._detect_chunk_type(paragraph)

            # If paragraph is too large, split further
            if len(paragraph) > self.max_chunk_size:
                sub_chunks = self._split_large_paragraph(paragraph)
                for sub_chunk in sub_chunks:
                    chunk = ChunkMetadata(
                        chunk_id=f"{document_id}_c{chunk_index}",
                        document_id=document_id,
                        document_version=document_version,
                        chunk_index=chunk_index,
                        chunk_type=chunk_type,
                        content=sub_chunk,
                        char_count=len(sub_chunk),
                        token_count=self._estimate_tokens(sub_chunk),
                        metadata=metadata or {},
                    )
                    chunks.append(chunk)
                    chunk_index += 1
            else:
                # Create chunk
                chunk = ChunkMetadata(
                    chunk_id=f"{document_id}_c{chunk_index}",
                    document_id=document_id,
                    document_version=document_version,
                    chunk_index=chunk_index,
                    chunk_type=chunk_type,
                    content=paragraph,
                    char_count=len(paragraph),
                    token_count=self._estimate_tokens(paragraph),
                    metadata=metadata or {},
                )
                chunks.append(chunk)
                chunk_index += 1

        return chunks

    def _split_paragraphs(self, content: str) -> List[str]:
        """Split content into paragraphs"""
        # Split by double newline or more
        paragraphs = re.split(r"\n\s*\n", content)
        return [p.strip() for p in paragraphs if p.strip()]

    def _split_large_paragraph(self, paragraph: str) -> List[str]:
        """Split large paragraph into smaller chunks"""
        chunks = []

        # Split by sentences
        sentences = re.split(r"(?<=[.!?])\s+", paragraph)

        current_chunk = ""
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.max_chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def _detect_chunk_type(self, text: str) -> ChunkType:
        """Detect chunk type from content"""
        # Heading detection (markdown style)
        if re.match(r"^#{1,6}\s+", text):
            return ChunkType.HEADING

        # List detection
        if re.match(r"^\s*[-*•]\s+", text, re.MULTILINE):
            return ChunkType.LIST

        # Code detection
        if text.startswith("```") or text.startswith("    "):
            return ChunkType.CODE

        # Table detection (simple heuristic)
        if "|" in text and text.count("|") > 2:
            return ChunkType.TABLE

        # Default to paragraph
        return ChunkType.PARAGRAPH

    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count (rough approximation)"""
        # Rough estimate: 1 token ≈ 4 characters
        return len(text) // 4


class DocumentProcessor:
    """
    Document Processor

    Unified processing pipeline for documents.
    """

    def __init__(self, chunker: Optional[Chunker] = None):
        self.chunker = chunker or Chunker()

    async def process_document(
        self,
        document_id: str,
        document_version: int,
        file_type: str,
        content: bytes,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[ChunkMetadata]:
        """
        Process document through pipeline

        Args:
            document_id: Document ID
            document_version: Document version
            file_type: MIME type
            content: File content
            metadata: Additional metadata

        Returns:
            List[ChunkMetadata]: Processed chunks
        """
        # Parse document
        text = await self._parse_document(file_type, content)

        # Clean text
        text = self._clean_text(text)

        # Normalize text
        text = self._normalize_text(text)

        # Chunk text
        chunks = self.chunker.chunk_text(
            document_id=document_id,
            document_version=document_version,
            content=text,
            metadata=metadata,
        )

        return chunks

    async def _parse_document(self, file_type: str, content: bytes) -> str:
        """
        Parse document content

        Note: This is a basic implementation. In production, use proper
        libraries like PyPDF2, python-docx, openpyxl, etc.
        """
        if file_type == "text/plain":
            return content.decode("utf-8")

        elif file_type == "text/markdown":
            return content.decode("utf-8")

        elif file_type == "text/csv":
            return content.decode("utf-8")

        elif file_type == "application/pdf":
            try:
                import io

                from pypdf import PdfReader

                reader = PdfReader(io.BytesIO(content))
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
                result = "\n\n".join(pages)
                if not result.strip():
                    result = f"[PDF Content - {len(content)} bytes, no extractable text]"
                return result
            except Exception as e:
                return f"[PDF Parse Error: {e}]"

        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                import io

                from docx import Document

                doc = Document(io.BytesIO(content))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(" | ".join(cells))
                return "\n\n".join(paragraphs) if paragraphs else "[DOCX Content - no text found]"
            except Exception as e:
                return f"[DOCX Parse Error: {e}]"

        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            try:
                import io

                from openpyxl import load_workbook

                wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
                sheets = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c) if c is not None else "" for c in row]
                        line = "\t".join(cells).strip()
                        if line:
                            rows.append(line)
                    if rows:
                        sheets.append(f"--- Sheet: {sheet_name} ---")
                        sheets.extend(rows)
                wb.close()
                return "\n".join(sheets) if sheets else "[XLSX Content - no data found]"
            except Exception as e:
                return f"[XLSX Parse Error: {e}]"

        else:
            return f"[Unsupported file type: {file_type}]"

    def _clean_text(self, text: str) -> str:
        """Clean text content"""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove control characters
        text = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", text)

        return text.strip()

    def _normalize_text(self, text: str) -> str:
        """Normalize text content"""
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(""", "'").replace(""", "'")

        # Normalize dashes
        text = text.replace("–", "-").replace("—", "-")

        # Normalize ellipsis
        text = text.replace("…", "...")

        return text
